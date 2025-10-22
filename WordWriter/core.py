# coding=utf-8
"""WordWriter 核心类模块 (v4.0.0)

这个模块包含 WordWriter 的面向对象实现。

Author: pzweuj
Since: v4.0.0
"""

from typing import Dict, List, Optional, Any
from docx import Document
from docx.table import Table

# 导入现有的函数式 API（作为底层实现）
from .WordWriter import (
    search_tag,
    replace_paragraph_string,
    replace_text_box_string,
    insert_picture,
    fill_table,
    remove_ele,
)
from .constants import TagPrefix, SpecialValue, LogMessage


# ============================================================================
# 标签搜索器类
# ============================================================================

class TagSearcher:
    """标签搜索器
    
    负责在 Word 文档中搜索所有标签。
    
    Attributes:
        document: Word 文档对象
        
    Example:
        >>> from docx import Document
        >>> doc = Document("template.docx")
        >>> searcher = TagSearcher(doc)
        >>> tags = searcher.search_all()
        >>> print(tags.keys())
    """
    
    def __init__(self, document: Document):
        """初始化标签搜索器
        
        Args:
            document: Word 文档对象
        """
        self.document = document
        
    def search_all(self) -> Dict[str, List]:
        """搜索文档中的所有标签
        
        Returns:
            标签字典，格式为 {tag_name: [tag_info, ...]}
        """
        tag_dict = {}
        
        self._search_headers_footers(tag_dict)
        self._search_paragraphs(tag_dict)
        self._search_tables(tag_dict)
        self._search_textboxes(tag_dict)
        
        return tag_dict
        
    def _search_headers_footers(self, tag_dict: Dict[str, List]) -> None:
        """搜索页眉页脚中的标签
        
        Args:
            tag_dict: 标签字典
        """
        sections_list = []
        for section in self.document.sections:
            sections_list.extend([
                section.header,
                section.first_page_header,
                section.footer,
                section.first_page_footer
            ])
        
        for section_part in sections_list:
            search_tag(tag_dict, section_part.paragraphs)
            
    def _search_paragraphs(self, tag_dict: Dict[str, List]) -> None:
        """搜索段落中的标签
        
        Args:
            tag_dict: 标签字典
        """
        search_tag(tag_dict, self.document.paragraphs)
        
    def _search_tables(self, tag_dict: Dict[str, List]) -> None:
        """搜索表格中的标签
        
        Args:
            tag_dict: 标签字典
        """
        for table in self.document.tables:
            rows = table.rows
            for row_idx in range(len(rows)):
                cells = rows[row_idx].cells
                for col_idx in range(len(cells)):
                    cell = cells[col_idx]
                    
                    if TagPrefix.TAG_START in cell.text and TagPrefix.TAG_END in cell.text:
                        if TagPrefix.TABLE in cell.text:
                            # 表格标签
                            tag = (TagPrefix.TABLE + "-" + 
                                   cell.text.split(TagPrefix.TABLE + "-")[1].split(TagPrefix.TAG_END)[0] + 
                                   TagPrefix.TAG_END)
                            tag_dict.setdefault(tag, []).append([table, row_idx, col_idx])
                        else:
                            # 单元格中的字符串标签
                            search_tag(tag_dict, cell.paragraphs)
                            
    def _search_textboxes(self, tag_dict: Dict[str, List]) -> None:
        """搜索文本框中的标签
        
        Args:
            tag_dict: 标签字典
        """
        from .constants import XMLNamespace
        
        children = self.document.element.body.iter()
        for child in children:
            if child.tag.endswith((XMLNamespace.TAG_ALTERNATE_CONTENT, XMLNamespace.TAG_TEXTBOX)):
                for child_item in child.iter():
                    if child_item.tag.endswith((XMLNamespace.TAG_RUN, XMLNamespace.TAG_PARAGRAPH_PROPERTIES)):
                        if child_item.text is not None:
                            if TagPrefix.TEXTBOX in child_item.text and TagPrefix.TAG_END in child_item.text:
                                tag_dict.setdefault(child_item.text.strip(), []).append(child_item)


# ============================================================================
# 内容替换器类
# ============================================================================

class ContentReplacer:
    """内容替换器
    
    负责替换文档中的标签内容。
    
    Attributes:
        document: Word 文档对象
        tag_dict: 标签字典
        
    Example:
        >>> replacer = ContentReplacer(document, tag_dict)
        >>> replacer.replace_all({"#[title]#": "新标题"})
    """
    
    def __init__(self, document: Document, tag_dict: Dict[str, List]):
        """初始化内容替换器
        
        Args:
            document: Word 文档对象
            tag_dict: 标签字典
        """
        self.document = document
        self.tag_dict = tag_dict
        
    def replace_all(self, replace_dict: Dict[str, str], logs: bool = True) -> None:
        """替换所有标签
        
        Args:
            replace_dict: 替换字典 {tag: value}
            logs: 是否打印日志
        """
        for tag_key, value in replace_dict.items():
            if tag_key not in self.tag_dict:
                if logs:
                    print(LogMessage.MISSING_TAG + tag_key)
                continue
                
            if logs:
                print(LogMessage.FILLING_TAG + tag_key)
                
            self._replace_tag(tag_key, value)
            
    def _replace_tag(self, tag: str, value: str) -> None:
        """替换单个标签
        
        Args:
            tag: 标签名称
            value: 替换值
        """
        if TagPrefix.TABLE in tag:
            self._replace_table(tag, value)
        elif TagPrefix.TEXTBOX in tag:
            self._replace_textbox(tag, value)
        elif TagPrefix.IMAGE in tag or TagPrefix.TABLE_IMAGE in tag:
            self._replace_image(tag, value)
        else:
            self._replace_text(tag, value)
            
    def _replace_text(self, tag: str, value: str) -> None:
        """替换文本标签
        
        Args:
            tag: 标签名称
            value: 替换值
        """
        for tag_item in self.tag_dict[tag]:
            replace_paragraph_string(tag_item[1], value)
            
    def _replace_image(self, tag: str, value: str) -> None:
        """替换图片标签
        
        Args:
            tag: 标签名称
            value: 图片路径
        """
        for tag_item in self.tag_dict[tag]:
            insert_picture(tag_item[1], tag, value)
            
    def _replace_table(self, tag: str, value: str) -> None:
        """替换表格标签
        
        Args:
            tag: 标签名称
            value: 表格文件路径或特殊值
        """
        if value == SpecialValue.DELETE_TABLE:
            for tag_item in self.tag_dict[tag]:
                remove_ele(tag_item[0])
        else:
            for tag_item in self.tag_dict[tag]:
                fill_table(tag_item[0], tag_item[1], tag_item[2], value)
                
    def _replace_textbox(self, tag: str, value: str) -> None:
        """替换文本框标签
        
        Args:
            tag: 标签名称
            value: 替换值
        """
        for tag_item in self.tag_dict[tag]:
            replace_text_box_string(tag_item, value)


# ============================================================================
# WordWriter 主类
# ============================================================================

class WordWriter:
    """WordWriter 主类
    
    提供面向对象的 API 来处理 Word 模板。
    
    Attributes:
        template_path: 模板文件路径
        document: Word 文档对象
        tag_dict: 标签字典
        
    Example:
        >>> # 方式1: 链式调用
        >>> WordWriter("template.docx") \\
        ...     .replace({"#[title]#": "报告"}) \\
        ...     .save("output.docx")
        
        >>> # 方式2: 分步调用
        >>> writer = WordWriter("template.docx")
        >>> writer.load()
        >>> writer.replace({"#[title]#": "报告"})
        >>> writer.save("output.docx")
        
        >>> # 方式3: 一步完成
        >>> WordWriter.process("template.docx", "output.docx", 
        ...                     {"#[title]#": "报告"})
    """
    
    def __init__(self, template_path: str):
        """初始化 WordWriter
        
        Args:
            template_path: 模板文件路径
        """
        self.template_path = template_path
        self.document: Optional[Document] = None
        self.tag_dict: Dict[str, List] = {}
        self._loaded = False
        self._searcher: Optional[TagSearcher] = None
        self._replacer: Optional[ContentReplacer] = None
        
    def load(self) -> 'WordWriter':
        """加载模板文档
        
        Returns:
            self，支持链式调用
            
        Raises:
            FileNotFoundError: 模板文件不存在
        """
        import os
        if not os.path.exists(self.template_path):
            raise FileNotFoundError(f"模板文件不存在: {self.template_path}")
            
        self.document = Document(self.template_path)
        self._searcher = TagSearcher(self.document)
        self.tag_dict = self._searcher.search_all()
        self._replacer = ContentReplacer(self.document, self.tag_dict)
        self._loaded = True
        
        return self
        
    def replace(self, replace_dict: Dict[str, str], logs: bool = True) -> 'WordWriter':
        """替换标签
        
        Args:
            replace_dict: 替换字典 {tag: value}
            logs: 是否打印日志
            
        Returns:
            self，支持链式调用
            
        Raises:
            RuntimeError: 文档未加载
        """
        if not self._loaded:
            self.load()
            
        if self._replacer is None:
            raise RuntimeError("Replacer not initialized")
            
        self._replacer.replace_all(replace_dict, logs)
        return self
        
    def save(self, output_path: str) -> None:
        """保存文档
        
        Args:
            output_path: 输出文件路径
            
        Raises:
            RuntimeError: 文档未加载
        """
        if not self._loaded or self.document is None:
            raise RuntimeError("文档未加载，请先调用 load() 方法")
            
        self.document.save(output_path)
        
    def get_tags(self) -> List[str]:
        """获取所有找到的标签列表
        
        Returns:
            标签名称列表
            
        Raises:
            RuntimeError: 文档未加载
        """
        if not self._loaded:
            self.load()
            
        return list(self.tag_dict.keys())
        
    @classmethod
    def process(cls, template_path: str, output_path: str, 
                replace_dict: Dict[str, str], logs: bool = True) -> None:
        """一步完成模板处理（类方法）
        
        这是一个便捷方法，等同于旧的函数式 API。
        
        Args:
            template_path: 模板文件路径
            output_path: 输出文件路径
            replace_dict: 替换字典
            logs: 是否打印日志
            
        Example:
            >>> WordWriter.process("template.docx", "output.docx",
            ...                     {"#[title]#": "报告"})
        """
        cls(template_path).replace(replace_dict, logs).save(output_path)
        
    def __enter__(self) -> 'WordWriter':
        """上下文管理器入口"""
        self.load()
        return self
        
    def __exit__(self, exc_type, exc_val, exc_tb) -> None:
        """上下文管理器出口"""
        # 清理资源（如果需要）
        pass
        
    def __repr__(self) -> str:
        """字符串表示"""
        status = "loaded" if self._loaded else "not loaded"
        tags_count = len(self.tag_dict) if self._loaded else 0
        return f"<WordWriter(template='{self.template_path}', status='{status}', tags={tags_count})>"
