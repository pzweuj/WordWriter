# coding=utf-8
# pzw
# 20230925
# v2.13 更新表格单元格中段落的style保持
# v2.12 增加合并单元格的函数，因为这个功能比较常用
# v2.11 修复bug，当表格跨页时也保持底边的样式
# v2.10 保持表格最后一行的边框样式
# v2.9 识别#DELETETHISPARAGRAPH#来删除段落，同时适用于图片标签
# v2.8 识别#DELETETHISPARAGRAPH#来删除段落
# v2.7 行距保持
# v2.6 更新表格寻找tag的方式
# v2.5 识别#DELETETHISTABLE#来删除表格
# v2.4 换行符处理，识别\x0a
# v2.3 修复表格标签放在第3行或以后，不能正常替换的bug
# v2.2 兼容1.0版本的TBIMG tag
# v2.1 可选输出log

import os
import pandas as pd
from docx import Document
from docx.oxml.ns import qn as nsqn
from docx.oxml import OxmlElement
# from docx.oxml.ns import nsdecls
# from docx.oxml import parse_xml
# from docx.enum.dml import MSO_THEME_COLOR_INDEX
# from docx.opc.constants import RELATIONSHIP_TYPE
# from docx.oxml.shared import OxmlElement
# from docx.oxml.shared import qn

# 通用搜索循环
def searchTag(tagDict, paragraphs):
    for p in paragraphs:
        if "#[" in p.text and "]#" in p.text:
            for r in p.runs:
                if "#[" in r.text and "]#" in r.text:
                    try:
                        tagDict[r.text.strip()].append([p, r])
                    except:
                        tagDict[r.text.strip()] = []
                        tagDict[r.text.strip()].append([p, r])

# 建立各类tag字典
## 遍历模板，从模板中寻找完整的tag
def searchTemplateTag(document):
    tagDict = {}

    ### 页眉页脚
    sectionsList = []
    for s in document.sections:
        sectionsList.append(s.header)
        sectionsList.append(s.first_page_header)
        sectionsList.append(s.footer)
        sectionsList.append(s.first_page_footer)
    for sl in sectionsList:
        searchTag(tagDict, sl.paragraphs)

    ### 段落
    searchTag(tagDict, document.paragraphs)

    ### 表格
    tables = document.tables
    for t in tables:
        rows = t.rows
        for r in range(len(rows)):
            cells = rows[r].cells
            for c in range(len(cells)):
                cell = cells[c]
                if "#[" in cell.text and "]#" in cell.text:
                    if "#[TABLE" in cell.text and "]#" in cell.text:
                        tag = "#[TABLE-" + cell.text.split("#[TABLE-")[1].split("]#")[0] + "]#"
                        try:
                            tagDict[tag].append([t, r, c])
                        except:
                            tagDict[tag] = []
                            tagDict[tag].append([t, r, c])
                    else:
                        searchTag(tagDict, cell.paragraphs)

    ### 文本框
    children = document.element.body.iter()
    for child in children:
        if child.tag.endswith(("AlternateContent", "textbox")):
            for ci in child.iter():
                if ci.tag.endswith(("main}r", "main}pPr")):
                    if ci.text != None:
                        if "#[TX" in ci.text and "]#" in ci.text:
                            try:
                                tagDict[ci.text.strip()].append(ci)
                            except:
                                tagDict[ci.text.strip()] = []
                                tagDict[ci.text.strip()].append(ci)
    
    return tagDict


## 超链接
# 功能是在一个段落后增加超链接，未找到文本替换的方法
# 参考 https://stackoverflow.com/questions/47666642/adding-an-hyperlink-in-msword-by-using-python-docx
# def add_hyperlink(paragraph, text, url):
#     part = paragraph.part
#     r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
#     hyperlink = OxmlElement("w:hyperlink")
#     hyperlink.set(qn("r:id"), r_id, )
#     new_run = OxmlElement("w:r")
#     rPr = OxmlElement("w:rPr")
#     new_run.append(rPr)
#     new_run.text = text
#     hyperlink.append(new_run)
#     r = paragraph.add_run()
#     r._r.append(hyperlink)
#     r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
#     r.font.underline = True
#     return hyperlink

# 获得指定行号表格边框底线格式
def get_table_bottom_border_details(tableObj, row_index, cell_index):    
    # 获取表格的指定行号
    last_row = tableObj.rows[row_index]

    # 获取表格的边框格式
    tbl_borders = tableObj._tbl.tblPr.first_child_found_in("w:tblBorders")
    tbl_bottom_border = tbl_borders.find(nsqn("w:bottom"))
    tbl_border_details = {
        'size': tbl_bottom_border.get(nsqn('w:sz'), '0'),
        'color': tbl_bottom_border.get(nsqn('w:color'), 'auto'),
        'space': tbl_bottom_border.get(nsqn('w:space'), '0'),
        'val': tbl_bottom_border.get(nsqn('w:val'), 'single'),
    }

    # 获取指定行号中所有单元格的底线边框格式
    bottom_border_details = []
    for cell in last_row.cells[cell_index:]:
        # 获取单元格的底线边框格式
        tc_borders = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcBorders")
        if tc_borders == None:
            bottom_border = None
        else:
            bottom_border = tc_borders.find(nsqn("w:bottom"))

        # 默认空样式
        border_details = {
            'size': '0',
            'color': 'auto',
            'space': '0',
            'val': 'nil'
        }

        if bottom_border is not None:
            border_details = {
                'size': bottom_border.get(nsqn('w:sz'), '0'),
                'color': bottom_border.get(nsqn('w:color'), 'auto'),
                'space': bottom_border.get(nsqn('w:space'), '0'),
                'val': bottom_border.get(nsqn('w:val'), 'single'),
            }
        
        bottom_border_details.append(border_details)

    
    return bottom_border_details, tbl_border_details

# 设置单元格底线边框格式
def set_cell_bottom_border(cell, styleList):
    size = styleList["size"]
    color = styleList["color"]
    space = styleList["space"]
    border_type = styleList["val"]
    tcPr = cell._tc.get_or_add_tcPr()
    tc_borders = tcPr.first_child_found_in("w:tcBorders")
    if tc_borders == None:
        tc_borders = OxmlElement("w:tcBorders")
        tcPr.append(tc_borders)

    bottom_border = tc_borders.find(nsqn("w:bottom"))
    if bottom_border == None:
        bottom_border = OxmlElement("w:bottom")
    
    # 设置底线边框的属性
    bottom_border.set(nsqn('w:sz'), size)
    bottom_border.set(nsqn('w:color'), color)
    bottom_border.set(nsqn('w:space'), space)
    bottom_border.set(nsqn('w:val'), border_type)
    tc_borders.append(bottom_border)

# 设置表格的底线边框格式
def set_table_bottom_border(table, styleList):
    size = styleList["size"]
    color = styleList["color"]
    space = styleList["space"]
    border_type = styleList["val"]
    tblPr = table._tbl.tblPr
    tbl_borders = tblPr.first_child_found_in("w:tblBorders")
    if tbl_borders == None:
        tbl_borders = OxmlElement("w:tblBorders")
        tblPr.append(tbl_borders)
    bottom_border = tbl_borders.find(nsqn("w:bottom"))
    if bottom_border == None:
        bottom_border = OxmlElement("w:bottom")

    # 设置底线边框的属性
    bottom_border.set(nsqn('w:sz'), size)
    bottom_border.set(nsqn('w:color'), color)
    bottom_border.set(nsqn('w:space'), space)
    bottom_border.set(nsqn('w:val'), border_type)
    tbl_borders.append(bottom_border)

## 字符串替换，适用于表格单元格中的字符串/页眉页脚字符串/段落字符串
def replaceParagraphString(run, replaceString):
    run.text = replaceString
    if replaceString == "#DELETETHISPARAGRAPH#":
        paragraph = run._element.getparent()
        remove_ele(paragraph)

## 图片插入，适用于表格中的图片和段落中的图片
def insertPicture(run, tag, picturePath):
    if os.path.exists(picturePath):
        run.text = ""
        if "(" in tag and ")" in tag:
            width = int(tag.split("(")[1].split(",")[0])
            height = int(tag.split(")")[0].split(",")[1])
            run.add_picture(picturePath, width*100000, height*100000)
        else:
            run.add_picture(picturePath)
    else:
        if picturePath == "#DELETETHISPARAGRAPH#":
            paragraph = run._element.getparent()
            remove_ele(paragraph)
        else:
            run.text = picturePath

## 文本框中字符串替换，仅适合于文本框内字符串
def replaceTextBoxString(childList, replaceString):
    for c in childList:
        c.text = replaceString

## 表格插入，通过插入一个以tab分割的txt文件插入表格
### 表格初始化
def OriginTableReadyToFill(tableFile):
    table = pd.read_csv(tableFile, header=None, sep="\t", dtype=str)
    return table

### 删除表格行
def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)

### 表格插入
def fillTable(table, row_id, cell_id, insertTable):
    tableToFill = OriginTableReadyToFill(insertTable)
    rowToFill = tableToFill.shape[0]
    columnToFill = tableToFill.shape[1]

    # 格式刷
    cellList = table.rows[row_id].cells[cell_id:]
    styleList = []
    for cell in cellList:
        p0 = cell.paragraphs[0]
        lineSpacingRule = p0.paragraph_format.line_spacing
        spaceAfter = p0.paragraph_format.space_after
        r0 = p0.runs[0]
        font = r0.font
        styleList.append([cell.vertical_alignment, p0.style, p0.alignment, r0.bold, r0.italic, r0.underline, font.name, font.size, font.color.rgb, font.highlight_color, lineSpacingRule, spaceAfter])

    # 获得标签行及最后一行的底边样式
    tagBottomStyle, tagTableStyle = get_table_bottom_border_details(table, row_id, cell_id)
    lastLineBottomStyle, lastLineTableStyle = get_table_bottom_border_details(table, -1, cell_id)

    # 将当前的最后一行的底边样式先处理为正常格式
    currentLastLine = table.rows[-1].cells[cell_id:]
    if len(currentLastLine) != len(tagBottomStyle):
        if len(tagBottomStyle) != 0:
            set_cell_bottom_border(currentLastLine[c], tagBottomStyle[0])
    else:
        for c in range(len(currentLastLine)):
            set_cell_bottom_border(currentLastLine[c], tagBottomStyle[c])

    # 判断行数是否足够，不够就添加
    if len(table.rows) - row_id < rowToFill:
        addRowAmount = rowToFill - len(table.rows) + row_id
        i = 0
        while i < addRowAmount:
            table.add_row()
            i += 1

    # 填充内容
    start = 0
    run_row = row_id
    while row_id <= rowToFill + run_row - 1:
        for co in range(columnToFill):
            tc = table.cell(row_id, co + cell_id)
            tc.text = str(tableToFill.iloc[start, co]).replace("\\x0a", "\n")
            tc.vertical_alignment = styleList[co][0]
            tc.paragraphs[0].style = styleList[co][1]
            tc.paragraphs[0].alignment = styleList[co][2]
            tc.paragraphs[0].paragraph_format.line_spacing = styleList[co][10]
            tc.paragraphs[0].paragraph_format.space_after = styleList[co][11]
            r = tc.paragraphs[0].runs[0]
            r.bold = styleList[co][3]
            r.italic = styleList[co][4]
            r.underline = False if styleList[co][5] != True else True
            r.font.name = styleList[co][6]
            if not r._element.rPr.rFonts == None:
                r._element.rPr.rFonts.set(nsqn("w:eastAsia"), r.font.name)
            r.font.size = styleList[co][7]
            r.font.color.rgb = styleList[co][8]
            r.font.highlight_color = styleList[co][9]

        start += 1
        row_id += 1
    
    # 删除空行
    for row in table.rows:
        pString = ""
        for cell in row.cells:
            for p in cell.paragraphs:
                pString = pString + p.text
        if pString == "":
            remove_row(table, row)

    # 处理表格的边框底线样式
    set_table_bottom_border(table, lastLineTableStyle)

    # 处理此时最后一行的边框底线样式
    newLastLine = table.rows[-1].cells[cell_id:]
    if len(newLastLine) != len(lastLineBottomStyle):
        if len(lastLineBottomStyle) != 0:
            if (lastLineBottomStyle[0]["size"] == "0") and (lastLineBottomStyle[0]["color"] == "auto"):
                # 这种情况认为是默认状态，那就优先表格底线样式
                set_cell_bottom_border(newLastLine[c], lastLineTableStyle)
            else:
                set_cell_bottom_border(newLastLine[c], lastLineBottomStyle[0])
    else:
        for c in range(len(newLastLine)):
            if (lastLineBottomStyle[0]["size"] == "0") and (lastLineBottomStyle[0]["color"] == "auto"):
                set_cell_bottom_border(newLastLine[c], lastLineTableStyle)
            else:
                set_cell_bottom_border(newLastLine[c], lastLineBottomStyle[c])

### 删除元素
def remove_ele(ele):
    if str(type(ele)) == "<class 'docx.oxml.text.paragraph.CT_P'>":
        ele.getparent().remove(ele)
    else:
        parent = ele._element.getparent()
        parent.remove(ele._element)

# 函数合并
def WordWriter(inputDocx, outputDocx, replaceDict, logs=True):
    template = Document(inputDocx)
    templateTagDict = searchTemplateTag(template)
    for k in replaceDict:
        if not k in templateTagDict:
            if logs:
                print("【Missing Tag】 " + k)
        else:
            if logs:
                print("【Filling Tag】 " + k)
            if "#[TABLE" in k:
                if replaceDict[k] == "#DELETETHISTABLE#":
                    for i in templateTagDict[k]:
                        tableID = i[0]
                        remove_ele(tableID)
                else:
                    for i in templateTagDict[k]:
                        fillTable(i[0], i[1], i[2], replaceDict[k])
            elif "#[TX" in k:
                for i in templateTagDict[k]:
                    replaceTextBoxString(i, replaceDict[k])
            elif "#[IMAGE" in k or "#[TBIMG" in k:
                for i in templateTagDict[k]:
                    insertPicture(i[1], k, replaceDict[k])
            else:
                for i in templateTagDict[k]:
                    replaceParagraphString(i[1], replaceDict[k])
    template.save(outputDocx)

# 合并内容相同的行，这些行需要是排好序的
def MergeTableRow(tableObj, colIndex, remove_other_row_text=True):
    # 获得需要合并的行
    rowLen = len(tableObj.rows)
    mergeList = []
    nowText = ""
    mergeStartPoint = mergeEndPoint = 0
    for i in range(rowLen):
        currentText = tableObj.rows[i].cells[colIndex].text
        if currentText != nowText:
            if mergeEndPoint > mergeStartPoint:
                mergeList.append([mergeStartPoint, mergeEndPoint])
            mergeEndPoint = i - 1
            mergeStartPoint = i
            nowText = currentText
        else:
            mergeEndPoint = i
    if mergeEndPoint > mergeStartPoint:
        mergeList.append([mergeStartPoint, mergeEndPoint])

    # 合并
    for m in mergeList:
        if remove_other_row_text:
            for j in range(m[0], m[1] + 1):
                cell = tableObj.cell(j, colIndex)
                if j == m[0]:
                    pass
                else:
                    cell.text = ""
                    for p in cell.paragraphs:
                        p.clear()
        tableObj.cell(m[0], colIndex).merge(tableObj.cell(m[1], colIndex))

