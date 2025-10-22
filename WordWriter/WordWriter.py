# coding=utf-8
# pzw
# 20251022
# v3.4.0 第二阶段重构：代码优化，提取函数，统一命名
# v3.3.0 第一阶段重构：代码清理，提取常量，统一命名
# v3.2   更新tag寻找算法，提升tag的查询能力
# v3.1   修复bottom无法找到border的问题
# v3.0.3 修复部分bug
# v3.0   解决run不完整的问题

import os
from typing import Dict, List, Tuple, Optional, Any
import pandas as pd
from docx import Document
from docx.table import Table, _Row, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.oxml.ns import qn as nsqn
from docx.oxml import OxmlElement

# 导入常量
from .constants import (
    TagPrefix,
    SpecialValue,
    Conversion,
    DefaultBorder,
    XMLNamespace,
    LogMessage
)

# ============================================================================
# 标签搜索辅助函数
# ============================================================================

def _contains_tag_markers(text: str) -> bool:
    """检查文本是否包含标签标记
    
    Args:
        text: 要检查的文本
        
    Returns:
        如果包含标签开始和结束标记返回 True
    """
    return TagPrefix.TAG_START in text and TagPrefix.TAG_END in text


def _is_simple_tag(text: str) -> bool:
    """判断是否为简单标签（单个完整标签）
    
    Args:
        text: 要判断的文本
        
    Returns:
        如果是简单标签返回 True
    """
    return (text.count(TagPrefix.TAG_START) == 1 and 
            text.count(TagPrefix.TAG_END) == 1)


def _extract_tag_name(text: str) -> str:
    """从文本中提取标签名称
    
    Args:
        text: 包含标签的文本
        
    Returns:
        完整的标签名称，如 "#[name]#"
    """
    start_pos = text.find(TagPrefix.TAG_START)
    end_pos = text.find(TagPrefix.TAG_END, start_pos)
    if start_pos != -1 and end_pos != -1:
        return text[start_pos:end_pos + len(TagPrefix.TAG_END)]
    return ""


def _process_simple_tag(tag_dict: Dict[str, List], paragraph: Paragraph) -> None:
    """处理简单标签（单个完整标签）
    
    Args:
        tag_dict: 标签字典
        paragraph: 包含标签的段落
    """
    tag_name = _extract_tag_name(paragraph.text)
    if tag_name:
        run_list = list(paragraph.runs)
        tag_dict.setdefault(tag_name, []).append([paragraph, run_list])


def _process_complex_tag(tag_dict: Dict[str, List], paragraph: Paragraph) -> None:
    """处理复杂标签（多个标签或跨 run）
    
    Args:
        tag_dict: 标签字典
        paragraph: 包含标签的段落
    """
    tag_parts = []
    run_list = []
    
    for run in paragraph.runs:
        text = run.text
        
        # 检查这个 run 是否包含完整的标签
        if TagPrefix.TAG_START in text and TagPrefix.TAG_END in text:
            # 单个 run 中的完整标签
            tag_name = _extract_tag_name(text)
            if tag_name:
                tag_dict.setdefault(tag_name, []).append([paragraph, [run]])
            # 重置状态，继续寻找下一个标签
            tag_parts = []
            run_list = []
        elif TagPrefix.TAG_START in text:
            # 找到标签开头（跨 run 的情况）
            tag_parts = [text]
            run_list = [run]
        elif TagPrefix.TAG_END in text:
            # 找到标签结尾（跨 run 的情况）
            tag_parts.append(text)
            run_list.append(run)
            if tag_parts:
                tag_name = "".join(tag_parts)  # 使用 join 一次性拼接
                tag_dict.setdefault(tag_name, []).append([paragraph, run_list])
            tag_parts = []
            run_list = []
        elif tag_parts:
            # 标签中间部分（跨 run 的情况）
            tag_parts.append(text)
            run_list.append(run)


# ============================================================================
# 主要标签搜索函数
# ============================================================================

# 通用搜索循环
## 形成的是类似{tag1: [p, r1, r2, r3], tag2: [p, r1]}这样的字典
def search_tag(tag_dict: Dict[str, List], paragraphs: List[Paragraph]) -> None:
    """搜索段落中的标签
    
    遍历段落列表，查找所有符合格式的标签，并将其添加到标签字典中。
    支持单个run中的完整标签和跨多个run的标签。
    
    Args:
        tag_dict: 标签字典，用于存储找到的标签
        paragraphs: 要搜索的段落列表
        
    Note:
        标签格式: #[标签名]#
        结果字典格式: {tag_name: [[paragraph, [run1, run2, ...]]]}
    """
    for paragraph in paragraphs:
        if not _contains_tag_markers(paragraph.text):
            continue
            
        if _is_simple_tag(paragraph.text):
            _process_simple_tag(tag_dict, paragraph)
        else:
            _process_complex_tag(tag_dict, paragraph)

# 建立各类tag字典
## 遍历模板，从模板中寻找完整的tag
def search_template_tag(document: Document) -> Dict[str, List]:
    """搜索模板文档中的所有标签
    
    遍历文档的各个部分（段落、表格、页眉页脚、文本框），
    查找所有符合格式的标签，并返回标签字典。
    
    Args:
        document: python-docx 的 Document 对象
        
    Returns:
        标签字典，格式为:
        {
            "#[tag1]#": [[paragraph, [run1, run2, ...]], ...],
            "#[TABLE-name]#": [[table, row_idx, col_idx], ...],
            "#[TX-name]#": [xml_element, ...],
            ...
        }
        
    Example:
        >>> from docx import Document
        >>> doc = Document("template.docx")
        >>> tags = search_template_tag(doc)
        >>> print(tags.keys())
        dict_keys(['#[title]#', '#[TABLE-data]#'])
    """
    tag_dict = {}

    ### 页眉页脚
    sections_list = []
    for section in document.sections:
        sections_list.extend([section.header, section.first_page_header, section.footer, section.first_page_footer])
    for section_part in sections_list:
        search_tag(tag_dict, section_part.paragraphs)

    ### 段落
    search_tag(tag_dict, document.paragraphs)

    ### 表格
    for table in document.tables:
        rows = table.rows
        for row_idx in range(len(rows)):
            cells = rows[row_idx].cells
            for col_idx in range(len(cells)):
                cell = cells[col_idx]
                if TagPrefix.TAG_START in cell.text and TagPrefix.TAG_END in cell.text:
                    if TagPrefix.TABLE in cell.text and TagPrefix.TAG_END in cell.text:
                        tag = TagPrefix.TABLE + "-" + cell.text.split(TagPrefix.TABLE + "-")[1].split(TagPrefix.TAG_END)[0] + TagPrefix.TAG_END
                        tag_dict.setdefault(tag, []).append([table, row_idx, col_idx])
                    else:
                        # 单元格中的字符串tag
                        search_tag(tag_dict, cell.paragraphs)

    ### 文本框，仅支持整个文本框中的内容替换
    children = document.element.body.iter()
    for child in children:
        if child.tag.endswith((XMLNamespace.TAG_ALTERNATE_CONTENT, XMLNamespace.TAG_TEXTBOX)):
            for child_item in child.iter():
                if child_item.tag.endswith((XMLNamespace.TAG_RUN, XMLNamespace.TAG_PARAGRAPH_PROPERTIES)):
                    if child_item.text != None:
                        if TagPrefix.TEXTBOX in child_item.text and TagPrefix.TAG_END in child_item.text:
                            tag_dict.setdefault(child_item.text.strip(), []).append(child_item)
    
    return tag_dict

# 获得指定行号表格边框底线格式
def get_table_bottom_border_details(
    table_obj: Table, 
    row_index: int, 
    cell_index: int
) -> Tuple[List[Dict[str, str]], Dict[str, str]]:    
    # 获取表格的指定行号
    last_row = table_obj.rows[row_index]

    # 默认空样式
    ## val: single 实线；dashed 虚线；nil 隐藏
    default_border_details = DefaultBorder.get_style_dict()

    # 获取表格的边框格式
    table_borders = table_obj._tbl.tblPr.first_child_found_in("w:tblBorders")
    if table_borders is not None:
        table_bottom_border = table_borders.find(nsqn("w:bottom"))
        table_border_details = {
            'size': table_bottom_border.get(nsqn('w:sz'), '0'),
            'color': table_bottom_border.get(nsqn('w:color'), 'auto'),
            'space': table_bottom_border.get(nsqn('w:space'), '0'),
            'val': table_bottom_border.get(nsqn('w:val'), 'single'),
        }
    else:
        table_border_details = default_border_details

    # 获取指定行号中所有单元格的底线边框格式
    bottom_border_details = []
    for cell in last_row.cells[cell_index:]:
        # 获取单元格的底线边框格式
        cell_borders = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcBorders")
        bottom_border = cell_borders.find(nsqn("w:bottom")) if cell_borders != None else None
        if bottom_border is not None:
            border_details = {
                'size': bottom_border.get(nsqn('w:sz'), '0'),
                'color': bottom_border.get(nsqn('w:color'), 'auto'),
                'space': bottom_border.get(nsqn('w:space'), '0'),
                'val': bottom_border.get(nsqn('w:val'), 'single'),
            }
        else:
            border_details = default_border_details
        bottom_border_details.append(border_details)
    return bottom_border_details, table_border_details

# ============================================================================
# 统一边框处理函数
# ============================================================================

from enum import Enum

class BorderTarget(Enum):
    """边框目标类型枚举"""
    CELL = "cell"
    TABLE = "table"


def _get_or_create_borders(target: Any, target_type: BorderTarget) -> OxmlElement:
    """获取或创建边框元素
    
    Args:
        target: 单元格或表格对象
        target_type: 目标类型
        
    Returns:
        边框元素
    """
    if target_type == BorderTarget.CELL:
        tcPr = target._tc.get_or_add_tcPr()
        borders = tcPr.first_child_found_in("w:tcBorders")
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tcPr.append(borders)
    else:  # TABLE
        tblPr = target._tbl.tblPr
        borders = tblPr.first_child_found_in("w:tblBorders")
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            tblPr.append(borders)
    return borders


def _set_border_attributes(borders: OxmlElement, style_dict: Dict[str, str]) -> None:
    """设置边框属性
    
    Args:
        borders: 边框元素
        style_dict: 样式字典
    """
    bottom = borders.find(nsqn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    
    bottom.set(nsqn('w:sz'), style_dict["size"])
    bottom.set(nsqn('w:color'), style_dict["color"])
    bottom.set(nsqn('w:space'), style_dict["space"])
    bottom.set(nsqn('w:val'), style_dict["val"])


def set_bottom_border(target: Any, style_dict: Dict[str, str], target_type: BorderTarget) -> None:
    """统一的底边框设置函数
    
    Args:
        target: 单元格或表格对象
        style_dict: 样式字典
        target_type: 目标类型
    """
    borders = _get_or_create_borders(target, target_type)
    _set_border_attributes(borders, style_dict)


# ============================================================================
# 兼容性包装函数
# ============================================================================

# 设置单元格底线边框格式
def set_cell_bottom_border(cell: _Cell, styleList: Dict[str, str]) -> None:
    """设置单元格底边框（兼容包装器）
    
    Args:
        cell: 单元格对象
        styleList: 样式字典
    """
    set_bottom_border(cell, styleList, BorderTarget.CELL)

# 设置表格的底线边框格式
def set_table_bottom_border(table: Table, styleList: Dict[str, str]) -> None:
    """设置表格底边框（兼容包装器）
    
    Args:
        table: 表格对象
        styleList: 样式字典
    """
    set_bottom_border(table, styleList, BorderTarget.TABLE)

## 字符串替换，适用于表格单元格中的字符串/页眉页脚字符串/段落字符串
def replace_paragraph_string(run_list: List[Run], replace_string: str) -> None:
    run_list[0].text = replace_string
    for idx, run in enumerate(run_list):
        if idx != 0:
            run.clear()
    if replace_string == SpecialValue.DELETE_PARAGRAPH:
        paragraph = run_list[0]._element.getparent()
        remove_ele(paragraph)

## 图片插入，适用于表格中的图片和段落中的图片
def insert_picture(run_list: List[Run], tag: str, picture_path: str) -> None:
    if os.path.isfile(picture_path):
        for run in run_list:
            run.text = ""
        if "(" in tag and ")" in tag:
            width = int(tag.split("(")[1].split(",")[0])
            height = int(tag.split(")")[0].split(",")[1])
            run_list[0].add_picture(picture_path, width*Conversion.CM_TO_EMU, height*Conversion.CM_TO_EMU)
        else:
            run_list[0].add_picture(picture_path)
    else:
        if picture_path == SpecialValue.DELETE_PARAGRAPH:
            paragraph = run_list[0]._element.getparent()
            remove_ele(paragraph)
        else:
            run_list[0].text = picture_path
            for idx, run in enumerate(run_list):
                if idx != 0:
                    run.clear()

## 文本框中字符串替换，仅适合于文本框内字符串
def replace_text_box_string(child_list: List[Any], replace_string: str) -> None:
    for child in child_list:
        child.text = replace_string

## 表格插入，通过插入一个以tab分割的txt文件插入表格
### 表格初始化
def load_table_from_file(table_file: str) -> pd.DataFrame:
    """从tab分隔的文本文件加载表格数据
    
    Args:
        table_file: tab分隔的文本文件路径
        
    Returns:
        pandas DataFrame 对象
        
    Example:
        >>> df = load_table_from_file("data.txt")
        >>> print(df.shape)
        (10, 5)
    """
    table = pd.read_csv(table_file, header=None, sep="\t", dtype=str)
    return table

### 删除表格行
def remove_row(table: Table, row: _Row) -> None:
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)

### 获取表格格式
def table_style_list(table: Table, rowIdx: int, cellIdx: int) -> List[List[Any]]:
    cell_list = table.rows[rowIdx].cells[cellIdx:]
    style_list = []
    for cell in cell_list:
        p0 = cell.paragraphs[0]
        lineSpacingRule = p0.paragraph_format.line_spacing
        spaceAfter = p0.paragraph_format.space_after
        r0 = p0.runs[0]
        font = r0.font
        style_list.append([cell.vertical_alignment, p0.style, p0.alignment, r0.bold, r0.italic, r0.underline, font.name, font.size, font.color.rgb, font.highlight_color, lineSpacingRule, spaceAfter])
    return style_list

### 表格内容填充及调整格式
def fill_table_text_and_style(
    table: Table, 
    row_id: int, 
    fill_table_id: pd.DataFrame, 
    fill_cell_id: int, 
    fill_row_id: int, 
    fill_col_id: int, 
    style_list: List[List[Any]]
) -> None:
    start = 0
    run_row = row_id
    while row_id <= fill_row_id + run_row - 1:
        for co in range(fill_col_id):
            tc = table.cell(row_id, co + fill_cell_id)
            tc.text = str(fill_table_id.iloc[start, co]).replace("\\x0a", "\n")
            tc.vertical_alignment = style_list[co][0]
            tc.paragraphs[0].style = style_list[co][1]
            tc.paragraphs[0].alignment = style_list[co][2]
            tc.paragraphs[0].paragraph_format.line_spacing = style_list[co][10]
            tc.paragraphs[0].paragraph_format.space_after = style_list[co][11]
            r = tc.paragraphs[0].runs[0]
            r.bold = style_list[co][3]
            r.italic = style_list[co][4]
            r.underline = False if style_list[co][5] != True else True
            r.font.name = style_list[co][6]
            if not r._element.rPr.rFonts == None:
                r._element.rPr.rFonts.set(nsqn("w:eastAsia"), r.font.name)
            r.font.size = style_list[co][7]
            r.font.color.rgb = style_list[co][8]
            r.font.highlight_color = style_list[co][9]

        start += 1
        row_id += 1

# ============================================================================
# 表格填充辅助函数
# ============================================================================

def _ensure_table_rows(table: Table, start_row: int, required_rows: int) -> None:
    """确保表格有足够的行数
    
    Args:
        table: 表格对象
        start_row: 起始行索引
        required_rows: 需要的行数
    """
    current_rows = len(table.rows) - start_row
    if current_rows < required_rows:
        rows_to_add = required_rows - current_rows
        for _ in range(rows_to_add):
            table.add_row()


def _is_row_empty(row: _Row) -> bool:
    """判断表格行是否为空
    
    Args:
        row: 表格行对象
        
    Returns:
        如果行为空返回 True
    """
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            if paragraph.text.strip():
                return False
    return True


def _remove_empty_rows(table: Table) -> None:
    """删除表格中的所有空行（优化版）
    
    Args:
        table: 表格对象
    """
    # 先收集所有空行，避免在遍历时修改
    empty_rows = [row for row in table.rows if _is_row_empty(row)]
    
    # 删除收集到的空行
    for row in empty_rows:
        remove_row(table, row)


def _apply_border_to_cells(
    cells: List[_Cell], 
    border_styles: List[Dict[str, str]], 
    fallback_style: Dict[str, str]
) -> None:
    """应用边框样式到单元格列表
    
    Args:
        cells: 单元格列表
        border_styles: 边框样式列表
        fallback_style: 备用样式
    """
    if len(cells) != len(border_styles):
        # 如果长度不匹配，使用第一个样式或备用样式
        style_to_use = border_styles[0] if border_styles else fallback_style
        for cell in cells:
            set_cell_bottom_border(cell, style_to_use)
    else:
        # 长度匹配，逐个应用
        for cell, style in zip(cells, border_styles):
            set_cell_bottom_border(cell, style)


### 表格插入
def fill_table(table: Table, row_id: int, cell_id: int, insertTable: str) -> None:
    tableToFill = load_table_from_file(insertTable)
    rowToFill = tableToFill.shape[0]
    columnToFill = tableToFill.shape[1]

    # 格式刷
    styleList = table_style_list(table, row_id, cell_id)

    # 获得标签行及最后一行的底边样式
    tagBottomStyle, tagTableStyle = get_table_bottom_border_details(table, row_id, cell_id)
    lastLineBottomStyle, lastLineTableStyle = get_table_bottom_border_details(table, -1, cell_id)

    # 将当前的最后一行的底边样式先处理为正常格式
    current_last_line = table.rows[-1].cells[cell_id:]
    _apply_border_to_cells(current_last_line, tagBottomStyle, tagTableStyle)

    # 确保表格有足够的行数
    _ensure_table_rows(table, row_id, rowToFill)

    # 填充内容
    fill_table_text_and_style(table, row_id, tableToFill, cell_id, rowToFill, columnToFill, styleList)
    
    # 删除空行
    _remove_empty_rows(table)

    # 处理表格的边框底线样式
    set_table_bottom_border(table, lastLineTableStyle)

    # 处理此时最后一行的边框底线样式
    new_last_line = table.rows[-1].cells[cell_id:]
    _apply_border_to_cells(new_last_line, lastLineBottomStyle, lastLineTableStyle)

### 删除元素
def remove_ele(ele: Any) -> None:
    """删除文档元素
    
    Args:
        ele: 要删除的元素
    """
    from docx.oxml.text.paragraph import CT_P
    
    # 使用 isinstance 代替字符串比较，性能更好
    if isinstance(ele, CT_P):
        ele.getparent().remove(ele)
    else:
        parent = ele._element.getparent()
        parent.remove(ele._element)

# 函数合并
def word_writer(
    input_docx: str, 
    output_docx: str, 
    replace_dict: Dict[str, str], 
    logs: bool = True
) -> None:
    """替换 Word 模板中的标签并生成新文档
    
    这是 WordWriter 的主函数，用于处理 Word 模板文件中的各种标签，
    包括文本替换、图片插入、表格填充等功能。
    
    Args:
        input_docx: 输入的模板文件路径，必须是有效的 .docx 文件
        output_docx: 输出的文件路径
        replace_dict: 替换字典，键为标签名，值为替换内容
            - 文本标签: "#[标签名]#" -> "替换文本"
            - 图片标签: "#[IMAGE-名称-(宽,高)]#" -> "图片路径"
            - 表格标签: "#[TABLE-名称]#" -> "表格文件路径"
            - 文本框标签: "#[TX-名称]#" -> "替换文本"
        logs: 是否打印日志信息，默认为 True
        
    Example:
        >>> replace_dict = {
        ...     "#[title]#": "报告标题",
        ...     "#[IMAGE-logo-(10,10)]#": "logo.png",
        ...     "#[TABLE-data]#": "data.txt"
        ... }
        >>> word_writer("template.docx", "output.docx", replace_dict)
        【Filling Tag】 #[title]#
        【Filling Tag】 #[IMAGE-logo-(10,10)]#
        【Filling Tag】 #[TABLE-data]#
        
    Note:
        - 表格文件应为 tab 分隔的文本文件
        - 图片尺寸单位为厘米
        - 特殊值 "#DELETETHISPARAGRAPH#" 可用于删除段落
        - 特殊值 "#DELETETHISTABLE#" 可用于删除表格
        
    Since:
        v1.0.0
    """
    template = Document(input_docx)
    template_tag_dict = search_template_tag(template)

    for tag_key in replace_dict:
        if not tag_key in template_tag_dict:
            if logs:
                print(LogMessage.MISSING_TAG + tag_key)
        else:
            if logs:
                print(LogMessage.FILLING_TAG + tag_key)
            if TagPrefix.TABLE in tag_key:
                if replace_dict[tag_key] == SpecialValue.DELETE_TABLE:
                    for tag_item in template_tag_dict[tag_key]:
                        table_id = tag_item[0]
                        remove_ele(table_id)
                else:
                    for tag_item in template_tag_dict[tag_key]:
                        fill_table(tag_item[0], tag_item[1], tag_item[2], replace_dict[tag_key])
            elif TagPrefix.TEXTBOX in tag_key:
                for tag_item in template_tag_dict[tag_key]:
                    replace_text_box_string(tag_item, replace_dict[tag_key])
            elif TagPrefix.IMAGE in tag_key or TagPrefix.TABLE_IMAGE in tag_key:
                for tag_item in template_tag_dict[tag_key]:
                    insert_picture(tag_item[1], tag_key, replace_dict[tag_key])
            else:
                for tag_item in template_tag_dict[tag_key]:
                    replace_paragraph_string(tag_item[1], replace_dict[tag_key])
    template.save(output_docx)

# 合并内容相同的行，这些行需要是排好序的
def merge_table_row(
    tableObj: Table, 
    colIndex: int, 
    remove_other_row_text: bool = True
) -> None:
    """合并表格中内容相同的连续行
    
    根据指定列的内容，合并内容相同的连续行。
    注意：表格必须已按该列排序。
    
    Args:
        tableObj: 要处理的表格对象
        colIndex: 用于判断合并的列索引
        remove_other_row_text: 是否清除被合并行的文本，默认 True
        
    Example:
        >>> from docx import Document
        >>> doc = Document("test.docx")
        >>> table = doc.tables[0]
        >>> merge_table_row(table, 0)  # 按第一列合并
        
    Since:
        v1.0.0
    """
    # 获得需要合并的行
    rowLen = len(tableObj.rows)
    mergeList = []
    nowText = ""
    mergeStartPoint = mergeEndPoint = 0
    for i in range(rowLen):
        currentText = tableObj.rows[i].cells[colIndex].text
        if currentText != nowText:
            if mergeEndPoint > mergeStartPoint:
                mergeList.append([mergeStartPoint, mergeEndPoint])
            mergeEndPoint = i - 1
            mergeStartPoint = i
            nowText = currentText
        else:
            mergeEndPoint = i
    if mergeEndPoint > mergeStartPoint:
        mergeList.append([mergeStartPoint, mergeEndPoint])

    # 合并
    for m in mergeList:
        if remove_other_row_text:
            for j in range(m[0], m[1] + 1):
                cell = tableObj.cell(j, colIndex)
                if j != m[0]:
                    cell.text = ""
                    for p in cell.paragraphs:
                        p.clear()
        tableObj.cell(m[0], colIndex).merge(tableObj.cell(m[1], colIndex))


# ============================================================================
# 向后兼容别名（保持旧的驼峰命名可用）
# 这些别名将在未来版本中弃用，请使用新的 snake_case 命名
# ============================================================================

# 主要函数别名
searchTag = search_tag
searchTemplateTag = search_template_tag
replaceParagraphString = replace_paragraph_string
replaceTextBoxString = replace_text_box_string
insertPicture = insert_picture
fillTable = fill_table
OriginTableReadyToFill = load_table_from_file
WordWriter = word_writer
MergeTableRow = merge_table_row
